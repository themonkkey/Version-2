#!/usr/bin/env python3
r"""Engine: district-specialised SUGGESTED interventions for every district.

Distinct from build_recommendations.py, which carries the PRESCRIPTION layer
(verbatim, page-verified quotes from a district's Vision & Action Plan, Kurnool
only for now). This engine builds the SUGGESTION layer that sits under the same
"Where <district> stands" card for a sector the plan names no interventions for.

The suggestions are authored, not quoted, so they are clearly labelled as
suggestions in the UI and are grounded in three sources, in this order:

  DIAGNOSIS  the district's own numbers for the sector (share, LQ, growth, rank,
             pattern), computed by build_recommendations.diagnose from
             landing/assets/dist/<District>.json. Deterministic.

  TOOLKIT    the sector playbook in gva_playbook.json: pathways, actions,
             policies for the sector the bucket maps to.

  PLAN       excerpts pulled from the district's own Vision & Action Plan PDF
             (pdftotext) where the plan mentions the sector, so a suggestion can
             be tailored to what the district itself has flagged.

Pipeline (the three stages are separate so generation can be re-run cheaply):

  build_suggested_interventions.py --briefs
        Stage 1 (deterministic). Writes scripts/_suggest_briefs.json: for every
        (district, bucket) a grounding brief = diagnosis + toolkit + plan
        excerpts. This is the input the generation stage reads.

  <generation>
        Stage 2. A model reads each brief and writes 3-5 specialised
        interventions plus a one-line grounding note per (district, bucket) into
        scripts/suggested_authored.json. Driven by the workflow in this repo's
        session, or any model given the briefs; the schema is documented below.

  build_suggested_interventions.py            (assemble; default)
  build_suggested_interventions.py --check    (validate only, exit 1 on drift)
        Stage 3 (deterministic). Joins the authored interventions onto the
        diagnosis and writes landing/assets/suggested_interventions.json, the
        artifact the site loads. --check re-derives the diagnosis and verifies
        coverage without writing.

scripts/suggested_authored.json schema:
  { "<district_slug>": { "<bucket>": { "interventions": ["..", ".."],
                                       "grounding": "one line" } } }
"""
import json, os, re, subprocess, sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import build_recommendations as BR

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIST_DIR = os.path.join(ROOT, "landing", "assets", "dist")
PLAYBOOK = os.path.join(ROOT, "landing", "assets", "gva_playbook.json")
PLAN_DIR = os.path.join(ROOT, "corpus_files", "vision_documents", "district")
# District field docs pulled from the programme Drive: per-constituency WOOP
# analyses and sector-wise GVA field-visit notes. The richest, most current
# grounding where it exists (19 of 28 districts).
DRIVE_DIR = os.path.join(ROOT, "corpus_files", "district_docs")
TXT_CACHE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_plan_txt")
DOC_CACHE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_drive_txt")
BRIEFS = os.path.join(os.path.dirname(os.path.abspath(__file__)), "_suggest_briefs.json")
AUTHORED = os.path.join(os.path.dirname(os.path.abspath(__file__)), "suggested_authored.json")
OUT = os.path.join(ROOT, "landing", "assets", "suggested_interventions.json")


def slugify(name):
    return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")


def norm_key(s):
    return re.sub(r"[^a-z0-9]", "", s.lower())


# dist json filename stem -> vision-plan folder name. Only the ones that differ
# by more than case need spelling out; the rest match on a normalised key.
PLAN_DIR_OVERRIDE = {
    "Ananthapuramu": "Ananthapur",
    "Anakapalle": "Anakapalli",
    "Dr.B.R.Ambedkar_Konaseema": "Dr._B.R._Ambedkar_Konaseema",
    "Nandyal": "Nandyala",
    "Ntr": "NTR",
    "Sps_Nellore": "Spsr_Nellore",
    "Ysr_Kadapa": "YSR_Kadapa",
    "Alluri_Seetha_Rama_Raju": "Alluri_Seetharama_Raju",
}


def plan_dir_for(stem):
    """Folder under vision_documents/district for a dist/<stem>.json, or None
    when the district has no published plan yet (new districts)."""
    if not os.path.isdir(PLAN_DIR):
        return None
    if stem in PLAN_DIR_OVERRIDE:
        cand = PLAN_DIR_OVERRIDE[stem]
        return cand if os.path.isdir(os.path.join(PLAN_DIR, cand)) else None
    dirs = [d for d in os.listdir(PLAN_DIR) if os.path.isdir(os.path.join(PLAN_DIR, d))]
    nk = norm_key(stem)
    for d in dirs:
        if norm_key(d) == nk:
            return d
    return None


def plan_text(stem):
    """pdftotext of the district's plan, cached. '' when no plan exists."""
    os.makedirs(TXT_CACHE, exist_ok=True)
    cache = os.path.join(TXT_CACHE, stem + ".txt")
    if os.path.exists(cache):
        return open(cache, encoding="utf-8").read()
    d = plan_dir_for(stem)
    if not d:
        return ""
    folder = os.path.join(PLAN_DIR, d)
    pdfs = [f for f in os.listdir(folder) if f.lower().endswith(".pdf")]
    if not pdfs:
        return ""
    # Prefer the underscored spelling when a folder carries two copies.
    pdfs.sort(key=lambda f: (f.count(" "), -len(f)))
    src = os.path.join(folder, pdfs[0])
    out = subprocess.run(["pdftotext", "-layout", src, "-"],
                         capture_output=True).stdout.decode("utf-8", "replace")
    open(cache, "w", encoding="utf-8").write(out)
    return out


import glob

# dist json stem -> Drive district-docs folder, where they differ by more than
# case/spacing. The rest match on a normalised key.
DRIVE_DIR_OVERRIDE = {
    "Alluri_Seetha_Rama_Raju": "ASR",
    "Ysr_Kadapa": "Kadapa",
    "Sps_Nellore": "Nellore",
    "Ntr": "NTR ( Vijayawada District)",
    "Visakhapatnam": "Vizag",
}


def drive_dir_for(stem):
    if not os.path.isdir(DRIVE_DIR):
        return None
    if stem in DRIVE_DIR_OVERRIDE:
        cand = os.path.join(DRIVE_DIR, DRIVE_DIR_OVERRIDE[stem])
        return cand if os.path.isdir(cand) else None
    nk = norm_key(stem)
    for d in os.listdir(DRIVE_DIR):
        if os.path.isdir(os.path.join(DRIVE_DIR, d)) and norm_key(d) == nk:
            return os.path.join(DRIVE_DIR, d)
    return None


def docx_text(path):
    """Paragraphs and table cells of one .docx; the WOOP docs carry most of their
    content in tables, so cells are pulled too."""
    try:
        import docx
    except ImportError:
        return ""
    try:
        d = docx.Document(path)
    except Exception:
        return ""
    out = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                out.append(" | ".join(dict.fromkeys(cells)))
    return "\n".join(out)


def drive_text(stem):
    """Concatenated text of every .docx in the district's Drive folder, cached.
    '' when the district has no Drive docs (9 of 28)."""
    os.makedirs(DOC_CACHE, exist_ok=True)
    cache = os.path.join(DOC_CACHE, stem + ".txt")
    if os.path.exists(cache):
        return open(cache, encoding="utf-8").read()
    folder = drive_dir_for(stem)
    if not folder:
        return ""
    docs = sorted(glob.glob(os.path.join(folder, "**", "*.docx"), recursive=True))
    docs = [d for d in docs if not os.path.basename(d).startswith("~$")]
    text = "\n".join(docx_text(d) for d in docs)
    open(cache, "w", encoding="utf-8").write(text)
    return text


# Bucket -> the words that mark a paragraph in the plan as being about it. Kept
# generous: a false positive costs a stray sentence in a brief, a false negative
# costs the grounding, so the trade favours recall.
BUCKET_KEYWORDS = {
    "crops":         ["crop", "paddy", "rice", "agricultur", "cultivat", "farmer", "seed", "irrigat", "yield", "kharif", "rabi"],
    "horticulture":  ["horticultur", "mango", "banana", "vegetable", "fruit", "flower", "chilli", "onion", "cashew", "oil palm", "turmeric"],
    "livestock":     ["livestock", "dairy", "milk", "poultry", "egg", "cattle", "buffalo", "sheep", "goat", "veterinary", "animal husband", "meat", "wool"],
    "fisheries":     ["fish", "aqua", "shrimp", "prawn", "pond", "marine", "fisher", "seafood", "brackish"],
    "forestry":      ["forest", "timber", "bamboo", "logging", "afforest", "silvicultur"],
    "mining":        ["mining", "mineral", "quarry", "granite", "limestone", "sand", "barytes", "geolog"],
    "manufacturing": ["manufactur", "industr", "factory", "msme", "textile", "food processing", "cluster", "industrial park", "mega"],
    "construction":  ["construction", "building", "cement", "housing", "infrastructur", "road", "real estate develop"],
    "electricity":   ["electricity", "power", "energy", "solar", "renewable", "grid", "water supply", "distribution"],
    "trade_hotels":  ["trade", "hotel", "restaurant", "tourism", "retail", "market yard", "hospitality", "commerce"],
    "transport":     ["transport", "logistic", "storage", "warehous", "railway", "port", "road connectivity", "cold chain"],
    "communications": ["communicat", "telecom", "broadband", "internet", "digital connect", "fibre", "mobile tower"],
    "financial":     ["bank", "insurance", "credit", "finance", "loan", "shg", "self help", "financial inclusion"],
    "real_estate":   ["real estate", "dwelling", "housing", "rent", "property", "land value"],
    "public_admin":  ["public administration", "governance", "citizen service", "revenue depart", "e-governance", "grievance"],
    "other_services": ["services", "education", "health", "skill", "it ", "software", "professional", "sanitation"],
}


def sentences(text):
    text = re.sub(r"[ \t]+", " ", text)
    # Split on line breaks and sentence ends; the plans are laid out as bullets
    # and short lines as often as prose, so both are treated as unit boundaries.
    parts = re.split(r"(?<=[.:;])\s+|\n", text)
    return [p.strip() for p in parts if len(p.strip()) >= 30]


def excerpts_for(text, bucket, cap=6):
    """Up to `cap` distinct plan lines that mention the bucket's keywords, with
    the sector's own name weighted first so the excerpt is about the sector, not
    a passing mention."""
    if not text:
        return []
    kws = BUCKET_KEYWORDS.get(bucket, [])
    hits, seen = [], set()
    for s in sentences(text):
        low = s.lower()
        score = sum(1 for k in kws if k in low)
        if not score:
            continue
        key = re.sub(r"\W+", "", low)[:80]
        if key in seen:
            continue
        seen.add(key)
        # drop table-noise lines that are mostly digits / headers in caps runs
        letters = sum(c.isalpha() for c in s)
        if letters < 20:
            continue
        hits.append((score, len(s), s))
    hits.sort(key=lambda t: (-t[0], t[1]))
    out = [h[2] for h in hits[:cap]]
    return [re.sub(r"\s+", " ", o)[:320] for o in out]


def playbook_index():
    pb = json.load(open(PLAYBOOK, encoding="utf-8"))
    idx = {}
    for g in pb["groups"]:
        for s in g["sectors"]:
            idx[s["key"]] = s
    return idx


def sentence(name, label, share, state_share, lq, growth):
    edge = ("above" if lq > 1 else "below") + " the state"
    move = ("grew %.2f percent" % growth) if growth >= 0 else ("contracted %.2f percent" % abs(growth))
    return ("%s is %.2f percent of %s's own GDVA against %.2f percent for the state, "
            "a location quotient of %.2f, %s. It %s in the latest year."
            % (label, share, name, state_share, lq, edge, move))


def diagnosis_rows(dist, name):
    """Per bucket for one district: everything the 'Where <district> stands'
    card renders (the same figures build_recommendations produces for Kurnool)."""
    rows, state_total = BR.diagnose(dist)
    out = []
    for border, (bkey, label, names) in enumerate(BUCKETS_LOCAL):
        parts = [rows[n] for n in names if n in rows]
        if not parts:
            continue
        share = sum(p["pct_of_district"] for p in parts)
        value = sum(p["value"] for p in parts)
        st = sum(p["state_sector_total"] for p in parts)
        state_share = st / state_total * 100.0
        lq = share / state_share
        growths = [{"name": n, "growth": rows[n]["growth"], "rank": rows[n]["rank"]}
                   for n in names if n in rows]
        lead_growth = growths[0]["growth"]
        lead_rank = growths[0]["rank"]
        pattern = BR.classify(lq, lead_growth, share)
        sec_key, grp_key = BR.SUBSECTOR_TO_PLAYBOOK[names[0]]
        out.append({
            "bucket": bkey, "order": border, "label": label,
            "playbook_sector": sec_key, "playbook_group": grp_key,
            "pattern": pattern, "pattern_label": BR.PATTERN_LABEL[pattern],
            "sub_sectors": names,
            "per_sub_sector": growths,
            "value_rs_cr": BR.fmt(value),
            "share_of_district_pct": BR.fmt(share),
            "state_share_pct": BR.fmt(state_share),
            "lq": BR.fmt(lq),
            "growth_pct": BR.fmt(lead_growth),
            "rank_among_districts": lead_rank,
            "sentence": sentence(name, label, share, state_share, lq, lead_growth),
        })
    return out


BUCKETS_LOCAL = BR.BUCKETS


def districts():
    return sorted(f[:-5] for f in os.listdir(DIST_DIR) if f.endswith(".json"))


def load_dist(stem):
    return json.load(open(os.path.join(DIST_DIR, stem + ".json"), encoding="utf-8"))


def build_briefs():
    pb = playbook_index()
    out = {}
    for stem in districts():
        dist = load_dist(stem)
        name = dist.get("name") or stem.replace("_", " ")
        text = plan_text(stem)
        docs = drive_text(stem)
        secs = []
        for d in diagnosis_rows(dist, name):
            sk = d["playbook_sector"]
            tool = pb.get(sk, {})
            secs.append({
                "bucket": d["bucket"],
                "label": d["label"],
                "diagnosis": {k: d[k] for k in ("share_of_district_pct", "lq",
                                                "growth_pct", "rank_among_districts",
                                                "pattern", "pattern_label")},
                "toolkit": {
                    "pathways": tool.get("pathways", []),
                    "actions": tool.get("actions", []),
                    "policies": tool.get("policies", []),
                },
                "plan_excerpts": excerpts_for(text, d["bucket"]),
                "doc_excerpts": excerpts_for(docs, d["bucket"], cap=6),
            })
        out[slugify(name)] = {"name": name, "stem": stem,
                              "has_plan": bool(text), "has_docs": bool(docs),
                              "sectors": secs}
    json.dump(out, open(BRIEFS, "w", encoding="utf-8"),
              ensure_ascii=False, indent=1)
    withplan = sum(1 for v in out.values() if v["has_plan"])
    nsec = sum(len(v["sectors"]) for v in out.values())
    print("wrote %s: %d districts (%d with a plan), %d sector briefs"
          % (os.path.relpath(BRIEFS, ROOT), len(out), withplan, nsec))


def assemble(check=False):
    authored = json.load(open(AUTHORED, encoding="utf-8")) if os.path.exists(AUTHORED) else {}
    stems = districts()
    districts_total = len(stems)
    briefs = json.load(open(BRIEFS, encoding="utf-8")) if os.path.exists(BRIEFS) else {}
    has_plan = {slugify(v["name"]): v.get("has_plan", False) for v in briefs.values()}
    has_docs = {slugify(v["name"]): v.get("has_docs", False) for v in briefs.values()}

    out = {"_note": "Authored SUGGESTIONS, not plan quotes. Grounded per district "
                    "in its diagnosis, the sector playbook (toolkit), its Vision "
                    "& Action Plan, and its Drive field/WOOP docs where present. "
                    "Built by scripts/build_suggested_interventions.py.",
           "districts_total": districts_total,
           "districts": {}}
    missing = []
    for stem in stems:
        dist = load_dist(stem)
        name = dist.get("name") or stem.replace("_", " ")
        slug = slugify(name)
        a = authored.get(slug, {})
        rows = diagnosis_rows(dist, name)
        # rank the sectors by share so the card can carry a #rank badge; the
        # verbatim-quote layer (Kurnool) keeps its own priority rank in
        # recommendations.json and is overlaid by the site at render time.
        order_by_share = sorted(range(len(rows)),
                                key=lambda i: -rows[i]["share_of_district_pct"])
        rank_of = {i: pos + 1 for pos, i in enumerate(order_by_share)}
        recs = []
        for i, d in enumerate(rows):
            iv = a.get(d["bucket"], {})
            ivs = iv.get("interventions", [])
            if not ivs:
                missing.append((slug, d["bucket"]))
            recs.append({
                "id": slug + "-" + d["bucket"],
                "bucket": d["bucket"], "order": d["order"], "rank": rank_of[i],
                "label": d["label"],
                "playbook_sector": d["playbook_sector"],
                "playbook_group": d["playbook_group"],
                "pattern": d["pattern"], "pattern_label": d["pattern_label"],
                "diagnosis": {
                    "value_rs_cr": d["value_rs_cr"],
                    "share_of_district_pct": d["share_of_district_pct"],
                    "state_share_pct": d["state_share_pct"],
                    "lq": d["lq"], "growth_pct": d["growth_pct"],
                    "rank_among_districts": d["rank_among_districts"],
                    "per_sub_sector": d["per_sub_sector"],
                    "sentence": d["sentence"],
                },
                "suggested": ivs,
                "grounding": iv.get("grounding", ""),
            })
        out["districts"][slug] = {
            "name": name, "has_plan": has_plan.get(slug, False),
            "has_docs": has_docs.get(slug, False),
            "data_year": dist.get("data_year") or dist.get("year"),
            "districts_total": districts_total,
            "recs": recs,
        }

    blob = json.dumps(out, ensure_ascii=False)
    if "—" in blob:
        sys.exit("em dash in output; project style forbids it")

    nrec = sum(len(v["recs"]) for v in out["districts"].values())
    print("%d districts, %d sector recs, %d still without authored suggestions"
          % (len(out["districts"]), nrec, len(missing)))
    if missing[:8]:
        print("  e.g. missing:", ", ".join("%s/%s" % m for m in missing[:8]))

    if check:
        if missing:
            sys.exit("--check: %d sector recs have no authored suggestions"
                     % len(missing))
        print("--check: verified, nothing written")
        return
    json.dump(out, open(OUT, "w", encoding="utf-8"),
              ensure_ascii=False, indent=2)
    print("wrote", os.path.relpath(OUT, ROOT))


def main():
    if "--briefs" in sys.argv:
        build_briefs()
    else:
        assemble(check="--check" in sys.argv)


if __name__ == "__main__":
    main()
